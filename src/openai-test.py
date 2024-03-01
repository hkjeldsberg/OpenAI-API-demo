from docx import Document
from openai import OpenAI

client = OpenAI(
    # defaults to os.environ.get("OPENAI_API_KEY")
    # api_key="My API Key",
)


def transcribe_audio(audio_file_path):
    with open(audio_file_path, 'rb') as audio_file:
        transcription = client.audio.transcriptions.create(file=audio_file, model="whisper-1")
    return transcription.text


def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)


def analyze_transcript(transcription, query, model="gpt-3.5-turbo"):
    response = client.chat.completions.create(
        model=model,
        temperature=1.5,
        messages=[
            {
                "role": "system",
                "content": query
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )

    return response.choices[0].message.content


def get_queries():
    summary_query = """
    You are a highly skilled AI trained in language comprehension and summarization. 
    I would like you to read the following text and summarize it into a concise abstract paragraph. 
    Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. 
    Please avoid unnecessary details or tangential points.
    """
    action_query = """
    You are an AI expert in analyzing conversations and extracting action items. 
    Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. 
    These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely.
    """
    sentiment_query = """
    As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. 
    Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. 
    Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
    """
    key_point_query = """
    You are a proficient AI with a specialty in distilling information into key points. 
    Based on the following text, identify and list the main points that were discussed or brought up. 
    These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. 
    Your goal is to provide a list that someone could read to quickly understand what was talked about.
    """

    queries = {
        "summary": summary_query,
        "key_point": key_point_query,
        "action": action_query,
        "sentiment": sentiment_query
    }

    return queries


def create_summary(transcript, queries):
    abstract_summary = analyze_transcript(transcript, queries["summary"])
    key_points = analyze_transcript(transcript, queries["key_point"])
    action_items = analyze_transcript(transcript, queries["action"])
    sentiment = analyze_transcript(transcript, queries["sentiment"])

    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }


if __name__ == '__main__':
    audio_file_path = "audio/EarningsCall.wav"
    doc_filename = "summary.docx"

    queries = get_queries()
    transcript = transcribe_audio(audio_file_path)
    summary = create_summary(transcript, queries)
    save_as_docx(summary, doc_filename)
